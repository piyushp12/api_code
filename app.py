from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docxtpl import DocxTemplate
from docx import Document
import json
from io import BytesIO
import tempfile
import subprocess
import re
import sys
import os
import jinja2
import base64

app = Flask(__name__)
CORS(app)

def validate_template(template_path):
    doc = Document(template_path)
    template_content = "\n".join([p.text for p in doc.paragraphs])

    for_loops = re.findall(r'{%\s*for\b', template_content)
    end_for_loops = re.findall(r'{%\s*endfor\b', template_content)
    if len(for_loops) != len(end_for_loops):
        raise ValueError("Incorrect or Missing'{% for %}' loop in template.")

def convert_to(folder, source, timeout=None):
    args = [libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]
    process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
    filename = re.search(r'-> (.*?) using filter', process.stdout.decode())
    return filename.group(1) if filename else None

def libreoffice_exec():
    if sys.platform == 'darwin':  
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    elif sys.platform == 'win32': 
        return r'C:\Program Files\LibreOffice\program\soffice.exe' 
    return 'libreoffice'

def add_placeholders(doc, data, parent_key=None, processed_keys=None):
    """
    Recursively add placeholders to the DOCX document based on the JSON structure, avoiding duplicate keys.
    """
    if processed_keys is None:
        processed_keys = set() 

    if isinstance(data, dict):
        if parent_key and parent_key not in processed_keys:
            doc.add_paragraph(f'{{#{parent_key}}}')  
            processed_keys.add(parent_key)

        for key, value in data.items():
            if key in processed_keys:
                continue  

            if isinstance(value, (dict, list)):  
                add_placeholders(doc, value, key, processed_keys)
            else:  # Primitive value
                doc.add_paragraph(f'{{{key}}}')
                processed_keys.add(key)

        if parent_key:
            doc.add_paragraph(f'{{/{parent_key}}}')  
    elif isinstance(data, list):
        if parent_key and parent_key not in processed_keys:
            doc.add_paragraph(f'{{#{parent_key}}}')  
            processed_keys.add(parent_key)

        for item in data:
            add_placeholders(doc, item, None, processed_keys)

        if parent_key:
            doc.add_paragraph(f'{{/{parent_key}}}')  

def generate_document(template_file, data, doc_type):
    """
    Generate document based on the provided template or create a new one dynamically from JSON.
    """
    if template_file:
        validate_template(template_file)
        template = DocxTemplate(template_file)
        try:
            template.render(data)
        except jinja2.TemplateSyntaxError as e:
            print("===========>85",e.message)
            raise ValueError("Missing 'endfor' in template")
        except jinja2.UndefinedError as e:
            print("===========>87",e.message)
            raise ValueError(f"Template error: Undefined variable encountered - {e.message}")

        output = BytesIO()
        template.save(output)
        output.seek(0) 
    else:
        doc = Document()
        add_placeholders(doc, data)
        output = BytesIO()
        doc.save(output)
        output.seek(0)

    if doc_type.lower() == 'pdf':
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx_file:
            temp_docx_file.write(output.getvalue())
            temp_docx_path = temp_docx_file.name

        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_filename = convert_to(temp_dir, temp_docx_path)
            if pdf_filename:
                pdf_path = os.path.join(temp_dir, pdf_filename)
                with open(pdf_path, 'rb') as pdf_file:
                    pdf_io = BytesIO(pdf_file.read())
                pdf_io.seek(0)
                return pdf_io, 'application/pdf', 'output.pdf'
            else:
                raise Exception("PDF conversion failed.")
    else:
        return output, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'output.docx'

@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    try:
        template_file = None
        json_data = None

        if 'template' in request.files:
            template_file = request.files['template']
            print("template_file",template_file)
        elif request.is_json:
            json_data = request.get_json()
            base64_template = json_data.get('template')
            if base64_template:
                template_bytes = base64.b64decode(base64_template)
                template_file = BytesIO(template_bytes)

        if 'data' in request.files:
            json_data_file = request.files['data']
            data = json.load(json_data_file)
        elif json_data and 'data' in json_data:
            data = json_data['data']
        elif 'data' in request.form:
            data = json.loads(request.form['data'])
        else:
            return jsonify({"error": "JSON data (either as file or raw JSON in form data) is required."}), 400

        doc_type = request.form.get('doc_type') if not request.is_json else json_data.get('doc_type')
        if not doc_type:
            return jsonify({"error": "Document type is required."}), 400

        file_io, mimetype, filename = generate_document(template_file, data, doc_type)
        return send_file(file_io, as_attachment=True, download_name=filename, mimetype=mimetype)

    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)